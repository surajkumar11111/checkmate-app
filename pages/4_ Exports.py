import streamlit as st
import pandas as pd
from dbconnection import fetch_cheque_details, init_db_connection
import io
from fpdf import FPDF
from docx import Document
from streamlit_lottie import st_lottie  # type: ignore
import json

# initialize Database Connection
init_db_connection()

# session state for cheque data
if "cheque_data" not in st.session_state:
    st.session_state.cheque_data = pd.DataFrame()

if st.session_state.cheque_data.empty:
    records = fetch_cheque_details()
    if records:
        expected_columns = ["cheque_date", "account_number", "bank_name", "cheque_number", "payee_name", "amount", "uploaded_at"]
        if len(records[0]) == 8:
            expected_columns.append("status")
        st.session_state.cheque_data = pd.DataFrame(records, columns=expected_columns)

df = st.session_state.cheque_data
df_readable = df.rename(columns=lambda x: x.replace("_", " ").title())

st.subheader("Export Cheque Records")

# handle screen width detection
from streamlit_javascript import st_javascript
screen_width = st_javascript("window.innerWidth")
if screen_width:
    st.session_state["viewport_width"] = int(screen_width)
elif "viewport_width" not in st.session_state:
    st.session_state["viewport_width"] = 1200

hide_animations = st.session_state["viewport_width"] <= 480

st.markdown("""
    <style>
        div[data-testid="column"] {
            padding: 0px !important;
        }
    </style>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)

with col1:
    def load_lottiefile(filepath: str):
        with open(filepath, "r") as f:
            return json.load(f)
    if not hide_animations:
        first_anim_lottie = load_lottiefile("images/Animation-Exports.json")
        st_lottie(
            first_anim_lottie,
            speed=1,
            loop=True,
            quality="low",
            width=510,
            key="expo-anim"
        )

# Styling for download buttons
with col2:
    st.markdown("""
    <style>
        .stDownloadButton { margin-top: 15px !important; }
        .stDownloadButton button {
            width: 90% !important;
            height: 60px !important;
            background-color: #0cc789 !important;
            color: white !important;
            border-radius: 5px !important;
            border: none !important;
            font-size: 16px !important;
            margin-bottom: 8px;
        }
        .stDownloadButton button:hover { background-color: #a0d1c1 !important; }
    </style>
    """, unsafe_allow_html=True)

    if df_readable is not None and not df_readable.empty:
        # convert DataFrame to CSV
        def convert_df_to_csv(dataframe):
            return dataframe.to_csv(index=False).encode('utf-8')

        csv_data = convert_df_to_csv(df_readable)
        st.download_button("Download as CSV", data=csv_data, file_name="cheque_records.csv", mime="text/csv")

        # convert DataFrame to PDF (Fixed logic)
        def convert_df_to_pdf(dataframe):
            pdf = FPDF("L", "mm", "A4")
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=10)

            col_widths = [max(pdf.get_string_width(str(col)) + 10, 30) for col in dataframe.columns]
            row_height = 8
            max_page_height = 190 

            # write column headers
            for i, col in enumerate(dataframe.columns):
                pdf.cell(col_widths[i], row_height, col, border=1)
            pdf.ln()

            for _, row in dataframe.iterrows():
                row_data = [str(item) for item in row]

                # calculate max line height for the row
                max_lines = max(pdf.get_string_width(item) // col_widths[i] + 1 for i, item in enumerate(row_data))
                dynamic_row_height = row_height * max_lines

                # check if a new page is required
                if pdf.get_y() + dynamic_row_height > max_page_height:
                    pdf.add_page()
                    for i, col in enumerate(dataframe.columns):
                        pdf.cell(col_widths[i], row_height, col, border=1)
                    pdf.ln()

                # write row data
                y_start = pdf.get_y()
                for i, item in enumerate(row_data):
                    x_start = pdf.get_x()
                    pdf.multi_cell(col_widths[i], row_height, item, border=1, align="L")
                    pdf.set_xy(x_start + col_widths[i], y_start)  # Move to next column
                pdf.ln(dynamic_row_height)

            # save PDF to memory
            output = io.BytesIO()
            pdf_data = pdf.output(dest="S").encode("latin1")
            output.write(pdf_data)
            output.seek(0)
            return output

        pdf_data = convert_df_to_pdf(df_readable)
        st.download_button("Download as PDF", data=pdf_data, file_name="cheque_records.pdf", mime="application/pdf")

        # convert DataFrame to Excel
        def convert_df_to_excel(dataframe):
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                dataframe.to_excel(writer, index=False, sheet_name="Cheque Records")
            output.seek(0)
            return output

        excel_data = convert_df_to_excel(df_readable)
        st.download_button("Download as Excel", data=excel_data, file_name="cheque_records.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # convert DataFrame to DOCX
        def convert_df_to_docx(dataframe):
            doc = Document()
            doc.add_heading("Cheque Records", level=1)

            table = doc.add_table(rows=1, cols=len(dataframe.columns))
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(dataframe.columns):
                hdr_cells[i].text = col_name

            for _, row in dataframe.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row):
                    row_cells[i].text = str(item)

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            return output

        docx_data = convert_df_to_docx(df_readable)
        st.download_button("Download as DOCX", data=docx_data, file_name="cheque_records.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("No cheque records available for export.")
